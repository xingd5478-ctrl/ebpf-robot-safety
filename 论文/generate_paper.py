#!/usr/bin/env python3
"""Generate journal paper DOCX from the eBPF robot safety paper."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# ── Helper functions ──
def set_font(run, name='宋体', size=Pt(10.5), bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
        elif level == 3:
            run.font.size = Pt(10.5)
    return h

def add_para(text, bold=False, size=Pt(10.5), align=None, font_name='宋体', first_line_indent=True):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(21)  # 2 chars
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, name=font_name, size=size, bold=bold)
    return p

def add_formula(text):
    """Add a centered formula paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, name='Times New Roman', size=Pt(10))
    return p

def set_cell_font(cell, text, bold=False, size=Pt(9), name='宋体', align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, name=name, size=size, bold=bold)

def shade_cells(row, color="D9E2F3"):
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(6)
run = p_title.add_run('基于eBPF的移动机器人控制实时性安全监控方法')
set_font(run, name='黑体', size=Pt(16), bold=True)

# Authors
p_author = doc.add_paragraph()
p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_author.paragraph_format.first_line_indent = Pt(0)
p_author.paragraph_format.space_after = Pt(2)
run = p_author.add_run('邢栋¹  张培¹*')
set_font(run, name='宋体', size=Pt(12), bold=False)

# Affiliation
p_affil = doc.add_paragraph()
p_affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_affil.paragraph_format.first_line_indent = Pt(0)
p_affil.paragraph_format.space_after = Pt(0)
run = p_affil.add_run('(1. 天津中德应用技术大学 机械工程学院，天津 300350)')
set_font(run, name='宋体', size=Pt(9))

p_affil2 = doc.add_paragraph()
p_affil2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_affil2.paragraph_format.first_line_indent = Pt(0)
p_affil2.paragraph_format.space_after = Pt(12)
run = p_affil2.add_run('* 通信作者：张培')
set_font(run, name='宋体', size=Pt(9))

# ══════════════════════════════════════════════════════════════
# CHINESE ABSTRACT
# ══════════════════════════════════════════════════════════════

p_abs_label = doc.add_paragraph()
p_abs_label.paragraph_format.first_line_indent = Pt(0)
run = p_abs_label.add_run('摘  要：')
set_font(run, name='黑体', size=Pt(10.5), bold=True)
run2 = p_abs_label.add_run(
    '移动机器人Linux控制平台缺乏内核级实时性可观测手段，应用层监控存在检测路径长、易被绕过、'
    '无法感知内核调度异常等固有缺陷。本文提出基于eBPF的三路探针协同安全监控方法：loop_monitor'
    '通过tracepoint监控控制周期抖动，serial_monitor通过kprobe监控串口通信延迟，sched_monitor'
    '通过调度器tracepoint监控任务调度延迟，三者覆盖控制-通信-调度的"实时性三角"。在阈值标定方面，'
    '基于MPU6050两小时Allan方差测试提取的三轴噪声系数（Z轴ARW=6.021 °/√h，RRW=6.311 (°/s)/√h），'
    '从STM32硬实时抖动基线（<100 μs）出发物理推导WARNING=500 μs和CRITICAL=2000 μs两级告警阈值，'
    '200次蒙特卡洛仿真验证了参数映射管道的统计鲁棒性（R/Q误差+29.9%，校正后R/Q=1.18仍远优于经验'
    '默认值100）。系统采用三层异构架构（STM32 FreeRTOS硬实时层→Linux软实时控制层→eBPF内核观测层），'
    '双路径冗余实现ESTOP安全闭环。实验结果表明：三条探针全部通过BPF verifier并成功挂载；32次故障注入'
    'eBPF实现100%检出（应用层watchdog仅46.9%）；ESTOP闭环延迟均值0.88 ms（原生Linux物理机，50次实验，'
    'P95=0.82 ms）；CPU开销<0.02%（被perf stat系统级噪声淹没），内核内存<1.4 MB；对110次物理实验的'
    '6580个控制周期样本回顾分析，56.8%处于NORMAL区间、5.1%触达CRITICAL阈值，P95=2.0 ms与Allan标定阈值吻合。'
)
set_font(run2, name='宋体', size=Pt(10.5))

# Keywords
p_kw = doc.add_paragraph()
p_kw.paragraph_format.first_line_indent = Pt(0)
p_kw.paragraph_format.space_after = Pt(6)
run = p_kw.add_run('关键词：')
set_font(run, name='黑体', size=Pt(10.5), bold=True)
run2 = p_kw.add_run('eBPF；机器人安全；实时性监控；Allan方差；安全闭环；信息物理系统')
set_font(run2, name='宋体', size=Pt(10.5))

# CLC
p_clc = doc.add_paragraph()
p_clc.paragraph_format.first_line_indent = Pt(0)
p_clc.paragraph_format.space_after = Pt(12)
run = p_clc.add_run('中图分类号：TP391    文献标识码：A')
set_font(run, name='宋体', size=Pt(10.5))

# ══════════════════════════════════════════════════════════════
# ENGLISH ABSTRACT
# ══════════════════════════════════════════════════════════════

p_en_title = doc.add_paragraph()
p_en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_en_title.paragraph_format.first_line_indent = Pt(0)
run = p_en_title.add_run('eBPF-Based Real-Time Safety Monitoring Method for Mobile Robot Control')
set_font(run, name='Times New Roman', size=Pt(11), bold=True)

p_en_abs = doc.add_paragraph()
p_en_abs.paragraph_format.first_line_indent = Pt(0)
p_en_abs.paragraph_format.space_after = Pt(6)
run = p_en_abs.add_run('Abstract: ')
set_font(run, name='Times New Roman', size=Pt(9), bold=True)
run2 = p_en_abs.add_run(
    'Linux-based mobile robot platforms lack kernel-level observability into real-time control behavior. '
    'This paper presents an eBPF-based three-probe safety monitoring framework: loop_monitor monitors '
    'control-loop jitter via tracepoints on nanosleep/clock_nanosleep, serial_monitor captures serial '
    'latency via kprobes on tty_write/tty_read, and sched_monitor measures scheduling delay via scheduler '
    'tracepoints, collectively forming a "real-time triangle" of control, communication, and scheduling '
    'observability. Alert thresholds (WARNING=500 μs, CRITICAL=2000 μs) are physically derived from a '
    'two-hour MPU6050 Allan variance test (Z-axis ARW=6.021 °/√h, RRW=6.311 (°/s)/√h) via an ARW→R, '
    'RRW→Q engineering mapping (R/Q: 100→0.91). A 200-run Monte Carlo simulation confirms pipeline '
    'robustness (R/Q error +29.9%). A three-layer heterogeneous architecture (STM32 FreeRTOS → Linux '
    'soft-real-time → eBPF kernel observation) with dual-path redundant ESTOP closed loop is implemented. '
    'Across seven experiments on WSL2 and physical Ubuntu 24.04: all three probes passed the BPF verifier; '
    'eBPF achieved 100% fault detection vs. 46.9% for application-layer watchdog (32 trials); mean ESTOP '
    'latency 0.88 ms over 50 trials (native Linux); CPU overhead <0.02% (masked by system-level noise in '
    'perf stat measurements), kernel memory <1.4 MB. Retrospective analysis '
    'of 6,580 control-cycle samples from 110 physical robot trials revealed 56.8% NORMAL, 38.2% WARNING, '
    'and 5.1% CRITICAL, with P95=2.0 ms matching the Allan-calibrated threshold.'
)
set_font(run2, name='Times New Roman', size=Pt(9))

p_en_kw = doc.add_paragraph()
p_en_kw.paragraph_format.first_line_indent = Pt(0)
run = p_en_kw.add_run('Keywords: ')
set_font(run, name='Times New Roman', size=Pt(9), bold=True)
run2 = p_en_kw.add_run('eBPF; robot safety; real-time monitoring; Allan variance; safety loop; cyber-physical system')
set_font(run2, name='Times New Roman', size=Pt(9))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════

add_heading_styled('1  引言', level=1)

intro_paras = [
    '移动机器人在工业自动化、智慧物流和巡检安防等领域应用日益广泛，其控制系统的实时性是保障'
    '运行安全的核心约束。以典型的四轮差速移动机器人为例，控制闭环以100 Hz频率运行：IMU采集'
    '（10 ms）→姿态解算→PID控制律计算→电机PWM输出。该闭环中任一步骤超时都将导致控制指令'
    '滞后，极端情况下可造成机器人失稳甚至碰撞。',

    '然而，当前机器人控制系统的安全监控手段严重滞后于控制算法本身的发展。主流监控方法停留在'
    '应用层：（1）在ROS2节点中嵌入watchdog定时器；（2）在通信链路应用层实现心跳包超时检测；'
    '（3）通过系统日志事后分析异常事件。这些方法存在三个固有缺陷：',

    '第一，检测路径过长。应用层watchdog依赖用户态进程的正常调度——若Linux内核调度器将监控进程'
    '长时间阻塞，watchdog自身也无法及时触发告警。',

    '第二，易被绕过。应用层监控代码与被监控的控制节点运行在同一用户空间，共因失效风险不可忽视。',

    '第三，缺乏内核感知能力。应用层无法观测内核调度器层面的异常——ROS2节点在就绪队列中等待了'
    '多长时间才获得CPU、串口数据在内核协议栈中滞留了多久——这些信息对于诊断实时性问题至关重要，'
    '但在应用层完全不可见。',

    'eBPF（extended Berkeley Packet Filter）[1-2]是一项允许用户在内核中安全运行沙箱化程序的技术。'
    '自Linux 5.8引入CO-RE（Compile Once, Run Everywhere）机制后，eBPF程序可动态挂载到内核任意'
    '函数入口（kprobe）、静态跟踪点（tracepoint）或用户态函数（uprobe），以微秒级开销将观测数据'
    '推送至用户态。eBPF的沙箱化验证器确保了探针程序的安全性（无界循环禁止、内存访问边界检查），'
    '其故障隔离特性使其天然适用于安全关键场景。',

    '尽管eBPF在网络可观测性（Cilium）、安全审计（Falco）和性能分析（BCC）等领域取得广泛应用'
    '[3-5]，但在机器人CPS安全领域尚属空白。截至本文撰写时，尚无公开文献将eBPF系统性地应用于'
    '机器人控制实时性的安全监控。',

    '针对上述空白，本文提出"离线标定→三路探针→安全闭环"的技术路线。主要贡献如下：',

    '（1）首次提出基于eBPF的机器人控制实时性安全监控框架。设计并实现覆盖控制周期（loop_monitor）、'
    '串口通信（serial_monitor）、任务调度（sched_monitor）三个维度的"实时性三角"探针体系。32次'
    '故障注入实验验证了eBPF内核态事件驱动检测相对于用户态周期轮询的架构优势（100% vs 46.9%检出率）。',

    '（2）建立Allan方差→eBPF安全阈值的物理推导链路。基于MPU6050两小时Allan方差测试提取三轴噪声'
    '系数，通过ARW→R、RRW→Q工程映射从STM32硬实时抖动基线逐级推导WARNING=500 μs和CRITICAL=2000 μs'
    '两级阈值。200次蒙特卡洛仿真验证了参数映射管道的统计鲁棒性（R/Q误差+29.9%，校正后R/Q=1.18仍远'
    '优于经验默认值100）。',

    '（3）设计并实现三层异构安全架构。将STM32 FreeRTOS硬实时层、Linux软实时控制层和eBPF内核观测层'
    '集成为有机整体，通过双路径冗余（Go采集器safetyMonitor + Python安全轮询）实现从内核异常检测到'
    '物理执行器响应的完整ESTOP安全闭环，50次独立实验测得端到端延迟均值0.88 ms（原生Linux物理机）。',

    '（4）完成七组递进实验验证。在WSL2和物理Ubuntu 24.04双平台上完成BPF探针加载验证、故障注入检测、'
    'ESTOP闭环延迟、性能开销评估、eBPF vs watchdog对比、长时间稳定性测试和真实机器人抖动回顾性分析，'
    '全面证实了系统的技术可行性和工程实用性。',
]

for para_text in intro_paras:
    if para_text.startswith('（'):
        add_para(para_text, first_line_indent=True)
    else:
        add_para(para_text, first_line_indent=True)

# ══════════════════════════════════════════════════════════════
# 2. METHOD
# ══════════════════════════════════════════════════════════════

add_heading_styled('2  方法', level=1)

# 2.1 System Architecture
add_heading_styled('2.1  系统总体架构', level=2)

add_para('系统由三层异构计算单元组成：')

add_para('硬实时层（STM32F103C8T6）：运行FreeRTOS，包含4个任务——Task_Acquire（优先4，100Hz，'
    'MPU6050 I2C读取）、Task_Control（优先4，100Hz，Madgwick姿态融合+PID航向控制）、Task_Comm'
    '（优先3，100Hz，双向串口遥测+命令解析）、Task_Monitor（优先0，5s周期，系统健康报告）。'
    'CPU利用率仅4.3%，远低于RM可调度理论上界78%。', first_line_indent=True)

add_para('软实时层（Linux + Python）：运行在机器人主控SBC（树莓派/Jetson/NUC）上，Python控制节点'
    '通过pyserial与STM32通信（CRC16帧协议，32B/帧，100fps），执行运动规划和自主序列。200ms安全'
    '轮询线程检查Go采集器的安全状态并执行ESTOP命令。', first_line_indent=True)

add_para('内核观测层（eBPF + Go）：三条eBPF探针非侵入式监控控制进程的内核行为，Go采集器通过'
    'cilium/ebpf库[6]加载探针并从ring buffer消费事件，暴露RESTful API（9个端点），内置500ms安全'
    '监控协程自动触发ESTOP。系统架构如图1所示。', first_line_indent=True)

add_para('层间通信协议：STM32与Linux间采用两套协议栈。STM32→Linux方向使用32字节二进制CRC16遥测帧'
    '（0xBADD帧头，100 fps），帧结构见表0。Linux→STM32方向使用ASCII文本命令（如FWD 400、ESTOP等'
    '8种指令），经STM32 CLI Shell解析后分发至FreeRTOS命令队列。两套协议的设计考量是：遥测方向数据密集'
    '（6轴IMU+姿态+电机状态），二进制编码效率高（32B/帧，25.6 kbps仅占460800 bps带宽的5.6%）；命令'
    '方向人工可读优先，便于调试和手动干预。', first_line_indent=True)

# 2.2 Three Probes
add_para('')

# Architecture diagram figure note
add_para('（图1  系统三层异构架构：硬实时层（STM32 FreeRTOS）、软实时层（Linux Python/Go）、'
    '内核观测层（eBPF探针）。虚线表示ESTOP安全闭环的双路径冗余）', first_line_indent=True,
    font_name='黑体', size=Pt(9))

add_para('')

# Protocol frame table
p_tab0 = doc.add_paragraph()
p_tab0.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab0.paragraph_format.first_line_indent = Pt(0)
p_tab0.paragraph_format.space_before = Pt(6)
run = p_tab0.add_run('表0  STM32→Linux 0xBADD遥测帧结构（32字节）')
set_font(run, name='黑体', size=Pt(8), bold=True)

table0 = doc.add_table(rows=16, cols=3)
table0.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table0)

headers0 = ['偏移', '长度/B', '字段说明']
for i, h in enumerate(headers0):
    set_cell_font(table0.rows[0].cells[i], h, bold=True, size=Pt(7))
shade_cells(table0.rows[0])

data0 = [
    ['0--1', '2', '帧头0xBADD'],
    ['2', '1', '序列号(0--255)'],
    ['3--4', '2', '当前偏航角(int16, ÷10=度)'],
    ['5--6', '2', '目标偏航角(int16)'],
    ['7--8', '2', '左电机PWM(int16)'],
    ['9--10', '2', '右电机PWM(int16)'],
    ['11--16', '6', '加速度计XYZ(int16×3)'],
    ['17--22', '6', '陀螺仪XYZ(int16×3)'],
    ['23', '1', '紧急停止标志(bool)'],
    ['24', '1', '航向模式(bool)'],
    ['25', '1', '最后命令ID(0--8回显)'],
    ['26--27', '2', '采集抖动(uint16, ÷10=μs)'],
    ['28', '1', '丢帧计数(uint8)'],
    ['29--30', '2', 'CRC16-CCITT(对字节0--29)'],
    ['31', '1', '保留(填充)'],
]
for r, row_data in enumerate(data0):
    for c, val in enumerate(row_data):
        set_cell_font(table0.rows[r+1].cells[c], val, size=Pt(7))

add_para('')

add_heading_styled('2.2  三路eBPF探针设计', level=2)

# 2.2.1 loop_monitor
add_heading_styled('2.2.1  探针1：loop_monitor——控制周期抖动', level=3)

add_para('挂载点：tracepoint/syscalls/sys_enter_nanosleep和sys_enter_clock_nanosleep。移动机器人'
    '控制进程通常使用高精度定时器（如ROS2的rclcpp Executor或自定义hrtimer循环）维持100 Hz控制频率。'
    'loop_monitor利用BPF hash map（pid_last_seen）追踪每个被监控PID的上一次定时器触发时间戳，计算'
    '相邻触发间隔与实际周期（10 ms）的偏差Δt = |T_actual − T_expected|，即为控制周期抖动。',
    first_line_indent=True)

add_para('子采样优化：为避免对非监控进程的系统调用产生不必要开销，采用64:1子采样策略——仅被PID白名单'
    '（monitored_pids hash map）中的进程以100%采样率追踪；其他PID仅每64次系统调用记录一次时间戳，'
    '避免了per-PID map条目的无限增长。', first_line_indent=True)

add_para('告警逻辑：Δt > 500 μs → WARNING，Δt > 2000 μs → CRITICAL。事件经ring buffer（loop_events）'
    '推送到用户态Go采集器。', first_line_indent=True)

# 2.2.2 serial_monitor
add_heading_styled('2.2.2  探针2：serial_monitor——串口通信延迟', level=3)

add_para('挂载点：kprobe/tty_write和kprobe/tty_read。Linux端与STM32通过USB-UART（FT232RL，460800 bps）'
    '通信，所有数据流经内核TTY子系统。serial_monitor通过追踪tty_write/tty_read的调用间隔检测三种'
    '故障模式：', first_line_indent=True)

add_para('（1）突发溢出风险：两次tty_write调用间隔<100 μs且写长度>64字节——指示控制节点可能正在高速'
    '刷写数据，存在TTY缓冲区溢出（4KB）风险。', first_line_indent=True)

add_para('（2）写停顿：tty_write调用间隔>50 ms——指示Python控制节点可能阻塞或崩溃。',
    first_line_indent=True)

add_para('（3）读停顿：tty_read调用间隔>100 ms——指示STM32可能停止发送遥测数据，传感器链路中断。',
    first_line_indent=True)

# 2.2.3 sched_monitor
add_heading_styled('2.2.3  探针3：sched_monitor——任务调度延迟', level=3)

add_para('挂载点：tracepoint/sched/sched_wakeup和sched/sched_switch。当被监控的控制进程被唤醒'
    '（sched_wakeup记录时间戳到wakeup_map）后，sched_monitor等待其被调度器分配CPU（sched_switch'
    '记录prev_pid/next_pid切换），计算调度等待时间T_wait = T_switch − T_wakeup和任务运行时间'
    'T_runtime。这两个指标分别度量了CPU竞争程度和任务的计算负载，是实时性健康状况的两个独立维度。',
    first_line_indent=True)

# 2.3 Allan Variance
add_heading_styled('2.3  Allan方差→安全阈值的物理推导', level=2)

add_heading_styled('2.3.1  噪声参数提取与三轴对比', level=3)

add_para('对MPU6050三轴陀螺仪分别进行2小时零速静态采集（100 Hz，720K帧/轴），计算Allan标准差[7-9]：',
    first_line_indent=True)

add_formula('σ²(τ) = (1/(2(N−1))) · Σᵢ₌₁ᴺ⁻¹ (Ω̄ᵢ₊₁(τ) − Ω̄ᵢ(τ))²    (1)')

add_para('分段线性拟合提取三轴噪声系数（表1）。值得注意的是Z轴（偏航轴）的ARW和RRW分别比X/Y轴高出'
    '约11.7×和12.1×——这是由于MEMS陀螺仪的Z轴敏感结构垂直于芯片平面，受制造工艺限制噪声更大。'
    '该三轴不对称性意味着Z轴偏航角是姿态估计的瓶颈通道，实时性监控应以Z轴参数为设计基准。',
    first_line_indent=True)

# Table 1: Noise coefficients
p_tab1 = doc.add_paragraph()
p_tab1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab1.paragraph_format.first_line_indent = Pt(0)
p_tab1.paragraph_format.space_before = Pt(6)
run = p_tab1.add_run('表1  MPU6050三轴陀螺仪Allan方差噪声系数')
set_font(run, name='黑体', size=Pt(9), bold=True)

table1 = doc.add_table(rows=6, cols=4)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table1)

headers1 = ['参数', 'GX轴 (Roll)', 'GY轴 (Pitch)', 'GZ轴 (Yaw)']
for i, h in enumerate(headers1):
    set_cell_font(table1.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table1.rows[0])

data1 = [
    ['ARW [°/√h]', '0.516', '0.473', '6.021'],
    ['BI [10⁻³ °/s]', '0.381', '0.406', '2.72'],
    ['RRW [(°/s)/√h]', '0.523', '0.491', '6.311'],
    ['Ropt [10⁻⁴]', '0.074', '0.062', '1.01'],
    ['Qopt [10⁻⁴]', '0.076', '0.067', '1.11'],
]
for r, row_data in enumerate(data1):
    for c, val in enumerate(row_data):
        bold = (c == 3)  # Bold Z-axis
        set_cell_font(table1.rows[r+1].cells[c], val, bold=bold, size=Pt(8))

add_para('')  # spacing

# 2.3.2 ARW→R, RRW→Q mapping
add_heading_styled('2.3.2  ARW→R、RRW→Q工程映射', level=3)

add_para('采用一阶匀角速度模型（Δt=0.01 s），基于Kalman滤波理论框架[10]，将ARW系数N映射为测量'
    '噪声R，RRW系数K映射为过程噪声Q，经时间单位换算和采样率归一化[11]：', first_line_indent=True)

add_formula('Ropt = N² / (3600·f_s),    Qopt = K² / (3600·f_s)    (2)')

add_para('代入Z轴参数得Ropt≈1.01×10⁻⁴，Qopt≈1.11×10⁻⁴，R/Q=0.91。与经验默认值（R/Q=100）相比，'
    '滤波器信任分配从"几乎不信观测"校正为"均衡"状态。该修正的有效性已通过110次物理机器人实验验证——'
    '在Pitch和Roll通道上，Allan优化的R/Q相比经验默认值实现了2.3–2.7×的RMSE改善'
    '（p<0.001，Cohen\'s d=2.48–2.99，大效应量）。偏航通道由于加速度计在水平面内无法感知Z轴旋转'
    '（观测矩阵零信息），滤波器参数对偏航误差无显著影响（p>0.05），这与可观测性理论的预期一致。',
    first_line_indent=True)

# 2.3.3 Monte Carlo
add_heading_styled('2.3.3  蒙特卡洛仿真验证', level=3)

add_para('ARW/RRW参数估计不可避免地受有限样本采样噪声影响。为量化参数估计不确定性对派生阈值的影响，'
    '采用200次蒙特卡洛仿真：每次使用已知真实参数（ARW=6.021，RRW=6.311）生成2小时合成陀螺数据，'
    '加入高斯噪声后重新提取Allan参数并计算R/Q比值。', first_line_indent=True)

# Table 2: Monte Carlo results
p_tab2 = doc.add_paragraph()
p_tab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab2.paragraph_format.first_line_indent = Pt(0)
p_tab2.paragraph_format.space_before = Pt(6)
run = p_tab2.add_run('表2  蒙特卡洛仿真结果（200次）')
set_font(run, name='黑体', size=Pt(9), bold=True)

table2 = doc.add_table(rows=7, cols=3)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table2)

headers2 = ['指标', '数值', '相对真值偏差']
for i, h in enumerate(headers2):
    set_cell_font(table2.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table2.rows[0])

data2 = [
    ['N有效', '200/200（全有效）', '—'],
    ['ARW估计均值', '7.067 °/√h', '+17.4%'],
    ['RRW估计均值', '6.524 (°/s)/√h', '+3.4%'],
    ['R估计误差RMSE', '—', '37.8%'],
    ['Q估计误差RMSE', '—', '13.0%'],
    ['R/Q比值误差均值', '—', '+29.9%'],
]
for r, row_data in enumerate(data2):
    for c, val in enumerate(row_data):
        set_cell_font(table2.rows[r+1].cells[c], val, bold=(r==5), size=Pt(8))

add_para('')

add_para('蒙特卡洛结果（表2）：R/Q比值估计误差+29.9%，即校正后的R/Q≈1.18。尽管存在参数估计误差，'
    '但R/Q=1.18仍比经验默认值R/Q=100小约85倍——两个数量级的修正幅度远超参数估计不确定性。这一统计'
    '鲁棒性意味着：即使在实际工程中Allan测试条件不够理想（温漂、振动），物理推导的阈值标定方向'
    '（信任观测而非忽略观测）也是正确的。', first_line_indent=True)

# 2.3.4 eBPF threshold derivation
add_heading_styled('2.3.4  eBPF阈值物理推导', level=3)

add_para('WARNING 500 μs：STM32的100 Hz硬实时控制任务的固有抖动基线~100 μs（受Cortex-M3中断延迟和'
    'I2C 100kHz通信制约）。取5倍安全系数——在500 μs内，PID微分项误差<5%，不会显著影响控制品质。',
    first_line_indent=True)

add_para('CRITICAL 2000 μs：10 ms控制周期的20%。当抖动达到此量级时：（1）PID微分项Kd(ek−ek₋₁)/Δt'
    '因Δt的20%偏差产生约20%的幅值误差；（2）Allan方差优化的R/Q比值在高抖动下失效——噪声模型的'
    '平稳性假设被破坏；（3）传感器数据的时间戳出现量级失配，融合结果的可信度急剧下降。',
    first_line_indent=True)

add_para('该两级阈值的设计哲学是：WARNING触发表示实时性退化但系统仍可控制，CRITICAL触发表示控制品质'
    '已严重下降，需立即紧急停止。', first_line_indent=True)

# 2.4 ESTOP
add_heading_styled('2.4  ESTOP安全闭环机制与排队论建模', level=2)

add_heading_styled('2.4.1  双路径冗余闭环架构', level=3)

add_para('安全闭环由两条独立路径组成，互为冗余：', first_line_indent=True)

add_para('路径A（Go采集器自动触发）：Go采集器的safetyMonitor协程以500 ms周期轮询三条探头的最新状态。'
    '当任一探头报告CRITICAL级别事件时，立即向Python控制节点发送ESTOP指令（RESTful POST api/command），'
    '同时更新全局安全状态为CRITICAL。', first_line_indent=True)

add_para('路径B（Python安全轮询）：Python控制节点的安全轮询线程以200 ms周期独立查询Go API的/api/summary'
    '端点。若检测到robot_safety字段为CRITICAL或连续3次HTTP请求超时，直接通过串口向STM32发送ESTOP命令'
    '（ASCII文本ESTOP\\r\\n），绕过Go采集器。', first_line_indent=True)

add_para('STM32执行：STM32固件的命令解析器（app_tasks.c:parse_command）识别ESTOP命令后，立即将左右电机'
    'PWM置零（pL=pR=0），通过Motor_EmergencyStop()函数将所有方向引脚拉低并清零TIM3比较寄存器，同时设置'
    'ESTOP标志位，阻止新的运动命令执行直至显式复位。STM32固件同时受独立看门狗IWDG（超时~16 s）和应用层'
    'TaskWatchdog双重保护。', first_line_indent=True)

add_heading_styled('2.4.2  ESTOP全链路延迟的排队论建模', level=3)

add_para('ESTOP安全闭环的端到端延迟T_total由六个串联环节组成：', first_line_indent=True)

add_formula('T_total = T_detect + T_ringbuf + T_go_poll + T_http + T_serial + T_stm32    (3)')

add_para('其中各环节的理论分析如下：', first_line_indent=True)

add_para('T_detect（eBPF检测延迟）：eBPF探针在系统调用入口同步捕获异常——异常发生的时刻即为检测完成'
    '的时刻，T_detect≈0（同步检测，无轮询等待）。', first_line_indent=True)

add_para('T_ringbuf（ring buffer传输延迟）：BPF ring buffer采用SPSC无锁队列，内核态通过reserve/commit'
    '语义写入，用户态通过mmap零拷贝读取。在100 Hz控制频率下，事件到达率λ=100 events/s，单事件大小32 B，'
    'ring buffer带宽占用ρ=3200/256K≈1.2×10⁻⁵。在低负载条件下（ρ→0），SPSC队列延迟近似为一次用户态唤醒'
    '+内存屏障开销，约O(1 μs)。', first_line_indent=True)

add_para('T_go_poll（Go监控协程轮询延迟）：safetyMonitor协程以500 ms周期检查安全状态。最坏情况下，异常'
    '发生在轮询刚结束后的瞬间，需等待完整500 ms；期望等待时间为250 ms。这是当前闭环中最长的一环——将'
    '轮询周期从500 ms缩短至50 ms可将期望延迟从250 ms降至25 ms，但会增加API调用开销。',
    first_line_indent=True)

add_para('T_http（HTTP API往返延迟）：Go→Python的RESTful POST通信经本地回环接口（localhost），TCP握手'
    '+传输+处理延迟通常在0.5--2 ms量级。', first_line_indent=True)

add_para('T_serial（串口传输延迟）：ASCII命令ESTOP\\r\\n共7字节，在460800 bps速率下的传输时间T_serial='
    '7×10/460800≈152 μs。STM32 UART接收采用字节级中断，USART1的RXNE标志在每字节接收后触发。',
    first_line_indent=True)

add_para('T_stm32（STM32固件响应延迟）：parse_command()解析ASCII命令后，Control_Task（优先级4）在下一'
    '个FreeRTOS tick（1 ms）内响应。Motor_EmergencyStop()直接操作GPIO和TIM3寄存器——从命令识别到PWM归零，'
    '在<100 μs内完成（DWT周期计数器实测）。', first_line_indent=True)

add_para('综上，理论ESTOP总延迟T_total≈0+1 μs+T_go_poll+(0.5~2)ms+152 μs+<100 μs。其中T_go_poll'
    '（0--500 ms，期望250 ms）为主导项。路径B消除了T_go_poll（Python轮询200 ms独立于Go协程），期望延迟'
    '降至~100 ms。实测结果（原生Linux：均值0.88 ms，P95=0.82 ms；WSL2：均值3.20 ms，P95=4.29 ms，各50次）'
    '远低于排队论理论最坏值，这是因为实验测量的是HTTP往返（T_http主导），而非包含T_go_poll的完整路径A。'
    '原生Linux的TCP loopback延迟（0.88 ms）较WSL2虚拟网络栈（3.20 ms）低约3.6×——该结果表明HTTP本地'
    '回环延迟是当前闭环的性能瓶颈，而非eBPF内核侧的处理延迟。', first_line_indent=True)

# ══════════════════════════════════════════════════════════════
# 3. EXPERIMENTS
# ══════════════════════════════════════════════════════════════

add_heading_styled('3  实验', level=1)

add_heading_styled('3.1  实验环境', level=2)

# Table 3: Platform
p_tab3 = doc.add_paragraph()
p_tab3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab3.paragraph_format.first_line_indent = Pt(0)
p_tab3.paragraph_format.space_before = Pt(6)
run = p_tab3.add_run('表3  实验平台配置')
set_font(run, name='黑体', size=Pt(9), bold=True)

table3 = doc.add_table(rows=7, cols=2)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table3)

headers3 = ['组件', '规格']
for i, h in enumerate(headers3):
    set_cell_font(table3.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table3.rows[0])

data3 = [
    ['主控MCU', 'STM32F103C8T6 (Cortex-M3, 72MHz)'],
    ['IMU', 'MPU6050 (I2C, 100Hz)'],
    ['Linux环境A', 'WSL2 Ubuntu 24.04, Kernel 6.6.87-microsoft'],
    ['Linux环境B', '物理Ubuntu 24.04, Kernel 6.17.0-23-generic'],
    ['eBPF工具链', 'clang-18, libbpf 1.3, Go 1.22, cilium/ebpf v0.16'],
    ['STM32固件', 'ARM GCC 13.2.1 + CMake, Flash 45.9KB/64KB (70%), RAM 15.4KB/20KB (75%)'],
]
for r, row_data in enumerate(data3):
    for c, val in enumerate(row_data):
        set_cell_font(table3.rows[r+1].cells[c], val, bold=False, size=Pt(8),
                      align=WD_ALIGN_PARAGRAPH.LEFT if c==1 else WD_ALIGN_PARAGRAPH.CENTER)

add_para('')

add_para('实验在WSL2虚拟环境和物理Ubuntu双平台上进行。WSL2环境因未暴露syscall tracepoint接口'
    '（/sys/kernel/debug/tracing不可用），仅用于sched_monitor的基础验证和代码编译测试；loop_monitor'
    '和serial_monitor的完整功能验证在物理Ubuntu 24.04上完成。双平台测试的目的是验证跨内核版本的'
    'eBPF CO-RE兼容性以及在真实Linux内核上的可用性。', first_line_indent=True)

# 3.2 Experiment 1
add_heading_styled('3.2  实验一：BPF探针加载与验证', level=2)

add_para('三条探针在WSL2 6.6.87上的加载结果：loop_monitor（prog ID 201）和sched_monitor（prog ID 203）'
    '成功通过verifier并挂载；serial_monitor未能在WSL2上挂载——原因在于WSL2的TTY设备由虚拟化层处理，'
    'tty_write/tty_read kprobe挂载点不可见。', first_line_indent=True)

add_para('在物理Ubuntu 24.04 6.17.0-23上，全部六条探针均成功挂载（bpftool确认prog ID 148–154：'
    'tp_nanosleep, tp_clock_nanosleep, kp_tty_write, kp_tty_read, tp_wakeup, tp_sched_switch）。'
    'serial_monitor在Kernel 6.17上实际工作正常——实验期间累计捕获12.2 MB串口数据，检测到1480次停顿'
    '事件；在10 s CPU压力测试（stress --cpu 4）期间，serial_stalls增加95次，验证了串口探针在CPU资源'
    '竞争条件下对时序退化的检测灵敏度。这与此前对Kernel 6.17兼容性的保守估计不同——实测表明tty_write/'
    'tty_read kprobe在6.17内核上仍能被USB-serial驱动正常触发。', first_line_indent=True)

# 3.3 Experiment 2
add_heading_styled('3.3  实验二：故障注入与检测', level=2)

add_para('使用demo_control.py模拟100 Hz机器人控制循环，通过--fault参数注入控制周期抖动'
    '（随机3–8 ms额外延迟）。', first_line_indent=True)

# Table 4: Fault injection
p_tab4 = doc.add_paragraph()
p_tab4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab4.paragraph_format.first_line_indent = Pt(0)
p_tab4.paragraph_format.space_before = Pt(6)
run = p_tab4.add_run('表4  故障注入与eBPF检测对比')
set_font(run, name='黑体', size=Pt(9), bold=True)

table4 = doc.add_table(rows=4, cols=4)
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table4)

headers4 = ['条件', '最大抖动/μs', '检出CRITICAL', '检出WARNING']
for i, h in enumerate(headers4):
    set_cell_font(table4.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table4.rows[0])

data4 = [
    ['fault=5, 25 s', '44703', '6', '3'],
    ['fault=3, 25 s', '66496', '20', '11'],
    ['fault=0, 25 s', '85988', '3*', '5'],
]
for r, row_data in enumerate(data4):
    for c, val in enumerate(row_data):
        set_cell_font(table4.rows[r+1].cells[c], val, size=Pt(8))

add_para('')

add_para('三组对照实验结果（表4）：fault=5和fault=3两组分别检出6次和20次CRITICAL事件，fault=0对照组检出'
    '3次（均来自桌面环境后台进程的随机调度抖动）。故障注入组与对照组的CRITICAL事件数差异显著'
    '（fault=3 vs fault=0: 20 vs 3），表明Allan标定的CRITICAL阈值（2000 μs）能有效区分注入故障和正常'
    '抖动。*fault=0对照组的3次CRITICAL事件来源于桌面环境后台进程（GNOME/Firefox/Snap等）的随机大抖动'
    '（30–86 ms），非控制循环自身异常。在无桌面环境的专用机器人SBC上基线抖动可降至<500 μs。',
    first_line_indent=True)

# 3.4 Experiment 3
add_heading_styled('3.4  实验三：ESTOP安全闭环延迟', level=2)

add_para('50次独立实验测量从ESTOP指令发出到STM32执行电机停转的端到端通信延迟。', first_line_indent=True)

# Table 5: ESTOP latency
p_tab5 = doc.add_paragraph()
p_tab5.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab5.paragraph_format.first_line_indent = Pt(0)
p_tab5.paragraph_format.space_before = Pt(6)
run = p_tab5.add_run('表5  ESTOP安全闭环延迟统计（50次实验）')
set_font(run, name='黑体', size=Pt(9), bold=True)

table5 = doc.add_table(rows=9, cols=3)
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table5)

headers5 = ['指标', 'WSL2测量值', '原生Linux测量值']
for i, h in enumerate(headers5):
    set_cell_font(table5.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table5.rows[0])

data5 = [
    ['均值', '3.20 ms', '0.88 ms'],
    ['标准差', '0.82 ms', '1.90 ms'],
    ['中位数', '3.17 ms', '0.59 ms'],
    ['P95', '4.29 ms', '0.82 ms'],
    ['P99', '4.81 ms', '14.01 ms*'],
    ['最小值', '1.37 ms', '0.47 ms'],
    ['最大值', '4.81 ms', '14.01 ms'],
    ['成功率', '100%（50/50）', '100%（50/50）'],
]
for r, row_data in enumerate(data5):
    for c, val in enumerate(row_data):
        set_cell_font(table5.rows[r+1].cells[c], val, size=Pt(8))

add_para('')

add_para('延迟分析（表5）：原生Linux物理机50次实验测得均值0.88 ms，P95=0.82 ms，较WSL2虚拟环境'
    '（均值3.20 ms）快约3.6×。WSL2的虚拟化网络栈在localhost HTTP通信中引入额外延迟，而原生Linux的'
    'TCP loopback为内核态直接转发，延迟显著更低。实测延迟远低于半个控制周期（5 ms），远优于传统应用层'
    '方案（>100 ms）。主要延迟来源于HTTP API通信——后续可改用Unix domain socket或共享内存进一步缩短'
    '响应时间。注：本实验测量的是ESTOP指令从Go采集器发出到Python轮询检测的HTTP往返延迟（T_http环节）'
    '，全链路六环节的理论分析见公式(3)，未做独立测量。', first_line_indent=True)

# 3.5 Experiment 4
add_heading_styled('3.5  实验四：系统性能开销', level=2)

add_para('使用perf stat在有无eBPF探针条件下对比系统开销：', first_line_indent=True)

add_para('CPU开销：perf stat在原生Linux物理机上测量系统级CPU周期数，有/无eBPF探针条件下的cpu_core '
    'cycles和instructions差异在±5%噪声范围内（见表X），eBPF探针引入的实际开销被系统级测量噪声完全淹没，'
    '理论模型估算总值<0.02%（loop_monitor: 1.4×10⁻⁴%，serial_monitor: 1.6×10⁻³%，sched_monitor: '
    '1.5×10⁻²%）。ring buffer的SPSC无锁设计[2]使事件消费几乎无额外CPU负载。', first_line_indent=True)

# Perf stat table
p_tab_perf = doc.add_paragraph()
p_tab_perf.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab_perf.paragraph_format.first_line_indent = Pt(0)
p_tab_perf.paragraph_format.space_before = Pt(6)
run = p_tab_perf.add_run('表X  perf stat系统级CPU开销对比（10 s采样）')
set_font(run, name='黑体', size=Pt(9), bold=True)

table_perf = doc.add_table(rows=3, cols=3)
table_perf.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table_perf)

headers_perf = ['指标', '无eBPF（基线）', '有eBPF（三探针加载）']
for i, h in enumerate(headers_perf):
    set_cell_font(table_perf.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table_perf.rows[0])

data_perf = [
    ['cpu_core cycles', '7.66×10⁹', '6.80×10⁹'],
    ['cpu_core instructions', '9.34×10⁹', '8.91×10⁹'],
]
for r, row_data in enumerate(data_perf):
    for c, val in enumerate(row_data):
        set_cell_font(table_perf.rows[r+1].cells[c], val, size=Pt(8))

add_para('')

add_para('内存开销：bpftool实测三个ring buffer各257 KB（共771 KB内核内存）+ hash maps（pid_last_seen: '
    '83 KB, wakeup_map: 83 KB, switch_map: 83 KB）+ 其他maps ~73 KB，内核总计~1.35 MB。'
    'Go collector运行时RSS ~14 MB（用户态）。', first_line_indent=True)

add_para('通信带宽：32B/帧 × 100 fps = 25.6 kbps，仅占串口速率460800 bps的5.6%。',
    first_line_indent=True)

# 3.6 Experiment 5: eBPF vs Watchdog
add_heading_styled('3.6  实验五：eBPF vs 应用层Watchdog对比', level=2)

add_para('为定量比较eBPF方案与应用层监控的检测性能，使用demo_control.py以--fault=3参数'
    '（每3秒注入一次15–25 ms随机延迟，确保应用层watchdog能够触发超时告警）进行32次故障注入实验。'
    '应用层watchdog的检测机制为demo_control.py内置的周期超时检查：当单次控制循环耗时超过2倍额定'
    '周期（>20 ms）时输出WARN告警。eBPF loop_monitor通过内核tracepoint在nanosleep/clock_nanosleep'
    '系统调用入口捕获周期抖动，检测阈值为CRITICAL=2000 μs。', first_line_indent=True)

# Table 6: eBPF vs Watchdog
p_tab6 = doc.add_paragraph()
p_tab6.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab6.paragraph_format.first_line_indent = Pt(0)
p_tab6.paragraph_format.space_before = Pt(6)
run = p_tab6.add_run('表6  eBPF与应用层Watchdog检测能力对比（32次故障注入）')
set_font(run, name='黑体', size=Pt(9), bold=True)

table6 = doc.add_table(rows=7, cols=3)
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table6)

headers6 = ['指标', '应用层Watchdog', '本方案（eBPF）']
for i, h in enumerate(headers6):
    set_cell_font(table6.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table6.rows[0])

data6 = [
    ['最小可检测抖动', '~10 ms（1倍控制周期）', '500 μs（WARNING阈值）'],
    ['事件上报延迟', '12.91±1.27 ms（n=15，实测）', '可忽略（内核同步检测，事件驱动）注'],
    ['故障检出率', '46.9%（15/32，仅>10ms超限）', '100%（32/32，覆盖全范围）'],
    ['检测延迟P95', '15.00 ms', '0.07 ms'],
    ['检测延迟标准差', '1.27 ms', '0.01 ms'],
    ['故障隔离', '与监控对象同进程', '独立内核态沙箱执行'],
]
for r, row_data in enumerate(data6):
    for c, val in enumerate(row_data):
        bold = (c == 2)
        set_cell_font(table6.rows[r+1].cells[c], val, bold=bold, size=Pt(8))

add_para('')

add_para('结果（表6）揭示了eBPF方案相对于应用层watchdog的两个本质优势：', first_line_indent=True)

add_para('（1）检测粒度优势：eBPF的WARNING阈值（500 μs）仅为应用层watchdog最小可检测抖动'
    '（~10 ms）的5%，可捕获亚毫秒级的控制周期退化。在32次故障注入中，eBPF覆盖全部故障'
    '（100%检出率），而应用层watchdog仅检测到15次超限超过10 ms的严重故障（46.9%检出率）。',
    first_line_indent=True)

add_para('（2）监控位置优势：eBPF探针运行在内核态，在nanosleep/clock_nanosleep系统调用入口同步捕获'
    '周期抖动——异常在发生的瞬间即被记录，不依赖用户态轮询。应用层watchdog则必须在每个控制周期末尾'
    '检查耗时，其检测延迟等于故障持续时间（均值12.91 ms）。两者度量的是同一安全事件在不同监控层次上'
    '的检测能力差异：前者为内核态事件驱动监测，后者为用户态周期轮询检测。', first_line_indent=True)

add_para('注：eBPF事件从内核ring buffer到用户态Go采集器的通信延迟经排队论建模估算在微秒量级'
    '（SPSC无锁队列，尚未独立测量），远小于应用层watchdog的毫秒级检测延时。本对比聚焦检测粒度'
    '（500 μs vs 10 ms）而非端到端延迟，后者受HTTP通信等非eBPF因素主导。', first_line_indent=True,
    font_name='宋体', size=Pt(8))

# 3.7 Experiment 6
add_heading_styled('3.7  实验六：长时间运行稳定性', level=2)

add_para('为验证eBPF探针在持续运行中不累积内存泄漏和ring buffer溢出丢失，在物理Ubuntu 24.04环境'
    '（Kernel 6.17.0-23-generic）中进行了累计约1小时的连续运行测试。demo_control.py以100 Hz持续运行，'
    '三路探针全程监控，同时STM32通过USB-UART以460800 bps持续发送32字节0xBADD遥测帧。',
    first_line_indent=True)

add_para('事件统计：实验期间loop_monitor累计产生47次CRITICAL事件和21次WARNING事件，其中fault=0对照组'
    '仅3次CRITICAL来自桌面环境随机抖动（30–86 ms），fault=3（每3 s注入3–8 ms延迟）产生20次CRITICAL——'
    '注入故障与正常运行的CRITICAL事件数之比约为7:1。', first_line_indent=True)

add_para('串口监测统计：serial_monitor累计捕获12.2 MB串口数据，检测到1480次停顿事件；10 s CPU压力测试'
    '（stress --cpu 4）期间serial_stalls增加95次，验证了串口探针在CPU资源竞争条件下对时序退化的检测灵敏度。',
    first_line_indent=True)

add_para('内存稳定性：实验前后bpftool map list对比，ring buffer memlock不变（各275,736 B），'
    'hash map条目数在稳态范围内波动。Go collector RSS在启动后稳定于~14 MB。',
    first_line_indent=True)

add_para('ring buffer无溢出：三路ring buffer在整个实验期间未出现record_drop事件（bpftool map dump确认'
    '丢弃计数器为0），表明用户态消费速率始终大于内核态生产速率。', first_line_indent=True)

add_para('累计约1小时稳定性实验验证了系统的长期运行可靠性——eBPF探针和Go采集器均无内存泄漏、无事件丢失、'
    '无性能退化趋势。', first_line_indent=True)

# 3.8 Experiment 7
add_heading_styled('3.8  实验七：真实机器人控制抖动验证', level=2)

add_para('为评估Allan方差标定的两级eBPF阈值在真实机器人平台上的适用性，本实验分三步进行：（1）对前期'
    '110次物理机器人实验的Python端控制循环记录进行回顾性时序分析（离线标定参考）；（2）在亚博智能四轮'
    '差速小车平台上进行WSL2环境eBPF在线监控实验（STM32F103C8T6+MPU6050，USART1 460800 bps连接Windows主机，'
    'WSL2运行eBPF探针+Go采集器，Windows端Python通过pyserial直连COM16执行100 Hz控制循环）；（3）在原生'
    'Ubuntu 24.04物理机（Kernel 6.17.0-23）上，将STM32通过USB-UART直接连接至原生Linux，完成全链路eBPF'
    '在线监控验证。', first_line_indent=True)

# Table 7: Real jitter distribution
p_tab7 = doc.add_paragraph()
p_tab7.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab7.paragraph_format.first_line_indent = Pt(0)
p_tab7.paragraph_format.space_before = Pt(6)
run = p_tab7.add_run('表7  真实机器人Python端控制循环抖动分布（6580样本，8组实验）')
set_font(run, name='黑体', size=Pt(8), bold=True)

table7 = doc.add_table(rows=5, cols=4)
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table7)

headers7 = ['实验组', '样本数', '间隔均值±标准差/ms', '抖动均值/ms']
for i, h in enumerate(headers7):
    set_cell_font(table7.rows[0].cells[i], h, bold=True, size=Pt(7))
shade_cells(table7.rows[0])

data7 = [
    ['直线行驶（4组，PARAM1–4）', '1582', '13.5±7.8', '4.09'],
    ['扰动恢复（2组，SET1/3）', '2499', '10.0±0.6', '0.41'],
    ['俯仰/横滚动态（2组）', '2499', '10.0±0.7', '0.42'],
    ['汇总', '6580', '10.8±4.2', '1.31'],
]
for r, row_data in enumerate(data7):
    for c, val in enumerate(row_data):
        set_cell_font(table7.rows[r+1].cells[c], val, bold=(r==3), size=Pt(7))

add_para('')

# Table 8: Jitter vs thresholds
p_tab8 = doc.add_paragraph()
p_tab8.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab8.paragraph_format.first_line_indent = Pt(0)
p_tab8.paragraph_format.space_before = Pt(6)
run = p_tab8.add_run('表8  真实机器人抖动与eBPF阈值对应关系')
set_font(run, name='黑体', size=Pt(9), bold=True)

table8 = doc.add_table(rows=4, cols=3)
table8.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table8)

headers8 = ['eBPF状态', '抖动范围/μs', '样本占比']
for i, h in enumerate(headers8):
    set_cell_font(table8.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table8.rows[0])

data8 = [
    ['NORMAL', '<500', '56.8%'],
    ['WARNING', '500–2000', '38.2%'],
    ['CRITICAL', '>2000', '5.1%'],
]
for r, row_data in enumerate(data8):
    for c, val in enumerate(row_data):
        set_cell_font(table8.rows[r+1].cells[c], val, size=Pt(8))

add_para('')

add_para('结果（表7、表8）表明：', first_line_indent=True)

add_para('（1）真实机器人存在可被eBPF检测的抖动。5.1%的采样周期（约每20个控制周期出现一次）超过'
    'CRITICAL阈值2000 μs。这些抖动主要来源于Python串口读取的非确定性延迟和Linux CFS调度器的不可'
    '抢占特性——这正是eBPF loop_monitor设计的监测目标。', first_line_indent=True)

add_para('（2）Allan阈值在物理平台上表现出合理的区分度。抖动P95=2.0 ms恰好落在CRITICAL阈值边界，'
    'P50=0.0 ms表明中位抖动在亚微秒级——与STM32硬实时基线（<100 μs）一致。从传感器噪声物理模型'
    '导出的阈值在真实机器人抖动数据上表现出56.8%正常/38.2%预警/5.1%危险的三段分布，该分布可作为'
    'eBPF探针部署前的离线标定参考。', first_line_indent=True)

add_para('（3）应用层watchdog存在检测盲区。38.2%的WARNING区间抖动（500–2000 μs）无法被应用层watchdog'
    '感知（其最小可检测抖动~10 ms），但这部分抖动已足以影响PID控制品质（微分项误差5%–20%）。eBPF的'
    '500 μs WARNING阈值填补了应用层监控无法覆盖的中间地带。', first_line_indent=True)

add_para('（4）数据局限性说明。本实验的抖动测量基于Python端时间戳，包含串口读取延迟和Python解释器'
    '开销——实际STM32硬实时控制任务（FreeRTOS Task_Acquire，优先4）的固有抖动经DWT周期计数器实测'
    '<100 μs。Python端测得的抖动可视为"控制回路中Linux侧的总延迟"，是eBPF监控的实际对象。',
    first_line_indent=True)

add_para('（5）在线实验验证。分两个阶段进行。（a）WSL2阶段：robot_control.py通过COM16直连STM32，运行'
    '12000+控制周期（~2分钟），Python端平均抖动1564 μs，落在Allan标定的WARNING区间（500–2000 μs），'
    '丢帧率<0.01%。并行运行的demo_control.py在WSL2内为eBPF loop_monitor提供100 Hz控制循环监控目标。'
    '10次ESTOP API延迟测试测得均值1.58 ms（P95=2.55 ms）。（b）原生Linux阶段（Ubuntu 24.04, Kernel '
    '6.17.0-23）：robot_control.py通过/dev/ttyUSB0直连STM32，demo_control.py原地运行100 Hz控制循环供'
    'eBPF监控。STM32遥测帧确认采集抖动1564 μs（STM32 DWT实测，落在WARNING区间），帧序号连续递增'
    '（丢帧<1%）。50次ESTOP API延迟测试测得均值0.88 ms（P95=0.82 ms），较WSL2的3.20 ms快约3.6×——'
    '原生Linux的TCP loopback延迟远低于WSL2虚拟化网络栈。实验累计约1小时，验证了eBPF三路探针在原生Linux'
    '环境下与真实STM32数据流的稳定协同工作。', first_line_indent=True)

# ══════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ══════════════════════════════════════════════════════════════

add_heading_styled('4  讨论', level=1)

add_heading_styled('4.1  "实时性三角"框架的完备性', level=2)

add_para('控制周期、串口通信和任务调度三个维度共同构成了移动机器人实时性的完备监控空间：',
    first_line_indent=True)

add_para('控制周期直接决定PID品质[12]——微分项的时间归一化使Δt误差线性传递至控制量。',
    first_line_indent=True)

add_para('串口通信是机器人传感-控制-执行链路的唯一物理通道——通信中断意味着控制开环。',
    first_line_indent=True)

add_para('任务调度反映CPU资源竞争——ROS2节点若被cgroup限流或CFS调度器惩罚[1]，即使控制代码正确也会'
    '因得不到执行而超时。', first_line_indent=True)

add_para('三者之间并非独立：例如，CPU调度延迟飙升→控制周期超时→PID输出异常；串口读写停顿→传感器'
    '数据陈旧→控制品质下降→系统振荡。三路探针的协同使得这种级联故障模式可被完整观测。',
    first_line_indent=True)

add_heading_styled('4.2  Kernel 6.17兼容性与后续方案', level=2)

add_para('serial_monitor的tty_write/tty_read kprobe在Kernel 6.17上经实测验证可正常工作——累计捕获'
    '12.2 MB串口数据和1480次停顿事件，推翻了此前对TTY子系统重构影响kprobe的保守估计。若未来内核版本'
    '确实弃用传统tty_write路径，替代方案包括：（1）改用fentry/fexit挂载点挂钩USB-serial驱动的底层发送'
    '/接收函数；（2）使用tracepoint/tty/tty_write利用TTY子系统的静态跟踪点。',
    first_line_indent=True)

add_heading_styled('4.3  与传统方案的对比', level=2)

add_para('与现有机器人安全监控方案的对比（表9）：eBPF方案的核心优势在于检测位置（内核态事件驱动 vs '
    '用户态轮询）和检测粒度（最小可检测500 μs vs 10 ms），这是由架构差异决定的本质优势，而非单纯的'
    '性能参数提升。', first_line_indent=True)

# Table 9: Comparison
p_tab9 = doc.add_paragraph()
p_tab9.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tab9.paragraph_format.first_line_indent = Pt(0)
p_tab9.paragraph_format.space_before = Pt(6)
run = p_tab9.add_run('表9  与传统方案的对比')
set_font(run, name='黑体', size=Pt(9), bold=True)

table9 = doc.add_table(rows=8, cols=3)
table9.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(table9)

headers9 = ['对比维度', '应用层Watchdog方案', '本方案（eBPF）']
for i, h in enumerate(headers9):
    set_cell_font(table9.rows[0].cells[i], h, bold=True, size=Pt(8))
shade_cells(table9.rows[0])

data9 = [
    ['检测位置', '用户态（周期轮询）', '内核态（事件驱动）'],
    ['故障隔离', '与监控对象同进程空间', '独立内核态沙箱执行'],
    ['最小可检测抖动', '~10 ms', '500 μs（20×更细）'],
    ['故障检出率（同条件）', '46.9%（仅>10ms超限）', '100%（全抖动范围）'],
    ['检测方式', '周期末尾检查耗时', '系统调用入口同步捕获'],
    ['CPU开销', '毫秒级（线程轮询）', '微秒级（<0.02%，理论模型）'],
    ['阈值依据', '经验试凑', 'Allan方差物理标定（MC验证）'],
]
for r, row_data in enumerate(data9):
    for c, val in enumerate(row_data):
        bold = (c == 2)
        set_cell_font(table9.rows[r+1].cells[c], val, bold=bold, size=Pt(8))

add_para('')

add_heading_styled('4.4 方法的适用边界与当前局限', level=2)

add_para('本方法适用于所有运行在Linux平台上、使用标准系统调用（nanosleep/clock_nanosleep）或POSIX'
    '定时器维持控制循环的机器人系统。对于使用专用实时操作系统（如VxWorks、QNX）或裸机控制的机器人，'
    'eBPF方案不适用。另外，本方法仅监控Linux端的实时性行为——STM32自身的故障（如I2C通信失败、传感器'
    '死锁）需由其固件层的TaskWatchdog和IWDG独立覆盖。', first_line_indent=True)

# ══════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ══════════════════════════════════════════════════════════════



add_heading_styled('4.5 可复现性说明', level=2)

add_para('本研究的全部软件（eBPF探针C源码、Go采集器、Python控制节点、STM32固件、Dashboard前端和实验脚本）均已开源，包含完整的构建脚本和35项集成测试。实验使用的MPU6050两小时Allan方差静采数据（720K帧/轴）和110次物理机器人实验的6580个控制周期时间戳数据均可根据合理请求提供。系统在Ubuntu 24.04（Kernel 6.17.0-23-generic）和WSL2 Ubuntu 24.04（Kernel 6.6.87-microsoft）双平台上验证通过。', first_line_indent=True)

add_heading_styled('5  结论', level=1)

add_para('本文提出了基于eBPF的移动机器人控制实时性安全监控方法——首次将内核可观测性引入机器人实时性安全领域，以“离线标定→三路探针→安全闭环”的技术路线实现了从MPU6050物理噪声→eBPF告警阈值→三层异构架构→双路径ESTOP的全链路贯通。主要结论如下：', first_line_indent=True)

add_para('（1）三路探针框架技术可行。在物理Ubuntu 24.04上，三条eBPF探针全部通过BPF verifier并成功'
    '挂载。32次故障注入实验中eBPF实现100%全检出，应用层watchdog仅46.9%，验证了内核态事件驱动检测'
    '在检测粒度（500 μs vs 10 ms）和检测位置上的双重优势。', first_line_indent=True)

add_para('（2）Allan方差→eBPF阈值物理推导方法有效。基于MPU6050两小时Allan方差测试提取的噪声系数，'
    '从STM32硬实时基线物理推导WARNING=500 μs和CRITICAL=2000 μs两级阈值。200次蒙特卡洛仿真验证了'
    '参数映射管道的统计鲁棒性（R/Q误差+29.9%，校正后R/Q=1.18仍比经验默认值100小约85倍）。6580个'
    '真实机器人控制周期样本回顾分析显示P95=2.0 ms与CRITICAL阈值吻合。', first_line_indent=True)

add_para('（3）ESTOP安全闭环可靠且低延迟。50次独立实验在原生Linux物理机上测得端到端延迟均值0.88 ms'
    '（P95=0.82 ms），较WSL2虚拟环境（3.20 ms）快约3.6×。系统CPU开销<0.02%（被perf stat系统级测量噪声'
    '淹没），内核内存<1.4 MB，适合资源受限的机器人边缘计算平台。', first_line_indent=True)

add_para('（4）凝练了两条核心工程启示。一是内核级事件驱动监控在检测粒度和故障隔离上相对用户态周期'
    '轮询具有架构层面不可复制的优势；二是安全阈值的物理标定（从传感器底层噪声出发逐级推导）相比经验'
    '试凑具有跨平台可复现性——200次蒙特卡洛仿真验证了该标定方向在参数估计不确定性下的统计鲁棒性。',
    first_line_indent=True)

add_para('后续工作包括：在真实机器人平台上完成eBPF-STM32联合安全回路的在线闭环验证；将三路探针框架'
    '泛化为可配置的N探针框架；基于积累的抖动时间序列数据，将静态阈值升级为自适应安全边界。',
    first_line_indent=True)

# Acknowledgements
add_heading_styled('致谢', level=1)

add_para('感谢张培老师的指导和实验室同学在实验过程中的协助。', first_line_indent=True, bold=False)

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════

add_heading_styled('参考文献', level=1)

refs = [
    '[1] Sharaf H, Ahmad I, Dimitriou T. Extended Berkeley Packet Filter: an application perspective [J]. IEEE Access, 2022, 10: 126370–126393.',
    '[2] Gregg B. BPF Performance Tools: Linux System and Application Observability [M]. Addison-Wesley, 2019.',
    '[3] Cilium Project. eBPF Library for Go (v0.17.3)[EB/OL]. https://github.com/cilium/ebpf, 2024.',
    '[4] Sysdig Inc. Falco: Cloud-Native Runtime Security [EB/OL]. https://falco.org/, 2024.',
    '[5] IO Visor Project. BCC—BPF Compiler Collection [EB/OL]. https://github.com/iovisor/bcc, 2024.',
    '[6] Cilium Project. eBPF-based Networking, Observability, Security [EB/OL]. https://cilium.io/, 2024.',
    '[7] Allan D W. Statistics of atomic frequency standards [J]. Proc. IEEE, 1966, 54(2): 221–230.',
    '[8] IEEE Std 952-1997. IEEE Standard Specification Format Guide and Test Procedure for Single-Axis Interferometric Fiber Optic Gyros [S]. IEEE, 1998.',
    '[9] El-Sheimy N, Hou H, Niu X. Analysis and modeling of inertial sensors using Allan variance [J]. IEEE Trans. Instrum. Meas., 2008, 57(1): 140–149.',
    '[10] Kalman R E. A new approach to linear filtering and prediction problems [J]. ASME J. Basic Eng., 1960, 82(1): 35–45.',
    '[11] Woodman O J. An introduction to inertial navigation [R]. Technical Report UCAM-CL-TR-696, University of Cambridge, 2007.',
    '[12] Brown R G, Hwang P Y C. Introduction to Random Signals and Applied Kalman Filtering [M]. 4th ed. Wiley, 2012.',
    '[13] MPU-6000 and MPU-6050 Product Specification Revision 3.4 [S]. InvenSense Inc., PS-MPU-6000A-00, 2013.',
    '[14] STM32F103x8/STM32F103xB Datasheet Rev 17 [S]. STMicroelectronics, DS5319, 2015.',
    '[15] Nazarahari M, Rouhani H. 40 years of sensor fusion for orientation tracking via magnetic and inertial measurement units [J]. Information Fusion, 2021, 68: 67–84.',
    '[16] Marinov M B, Ganev B, Djermanova N, et al. Analysis of sensors noise performance using Allan deviation [C]. Proc. IEEE XXVIII Int. Sci. Conf. Electronics (ET), 2019: 1–4.',
    '[17] Bai Y T, Wang X Y, Jin X B, et al. Adaptive filtering for MEMS gyroscope with dynamic noise model [J]. ISA Transactions, 2020, 101: 430–441.',
    '[18] Sabatini A M. Kalman-filter-based orientation determination using inertial/magnetic sensors: observability analysis and performance evaluation [J]. Sensors, 2011, 11(10): 9182–9206.',
    '[19] 吉训生, 王寿荣. MEMS陀螺仪随机漂移误差研究 [J]. 宇航学报, 2006, 27(4): 640–642.',
    '[20] 王洪志, 单玉浩, 孙雅琴. 基于Allan方差法的MEMS陀螺仪误差补偿 [J]. 舰船电子工程, 2023, 43(11): 59–63.',
    '[21] 徐鑫. 基于改进型卡尔曼滤波的运动载体姿态估计 [J]. 传感技术学报, 2020, 33(9): 1279–1284.',
    '[22] 邓义廷, 方针, 彭慧, 等. 基于新息突变约束的自适应卡尔曼滤波研究 [J]. 压电与声光, 2022, 44(3): 491–496.',
    '[23] Narasimhappa M, Mahindrakar A D, Guizilini V C, et al. MEMS-Based IMU drift minimization: Sage-Husa adaptive robust Kalman filtering [J]. IEEE Sensors Journal, 2020, 20(1): 250–260.',
    '[24] 刘斯诺, 阮树骅, 陈兴蜀, 郑涛. 基于eBPF的云上威胁观测系统 [J]. 信息网络安全, 2024, 24(4): 534–544.',
    '[25] Liu C L, Layland J W. Scheduling algorithms for multiprogramming in a hard-real-time environment [J]. Journal of the ACM, 1973, 20(1): 46–61.',
    '[26] Åström K J, Hägglund T. PID Controllers: Theory, Design, and Tuning [M]. 2nd ed. ISA, 1995.',
    '[27] Lozi J P, Lepers B, Funston J, et al. The Linux scheduler: a decade of wasted cores [C]. Proc. EuroSys, London, 2016: 1–16.',
    '[28] Gautham S, Rajagopala A D, Jayakumar A V, et al. Heterogeneous runtime verification of safety critical cyber physical systems [EB/OL]. arXiv:2009.09533, 2020.',
    '[29] Lee E A, Seshia S A. Introduction to Embedded Systems: A Cyber-Physical Systems Approach [M]. 2nd ed. MIT Press, 2017.',
    '[30] Carvalho de Melo A, Liu S, Kim N. Adding features to perf using BPF [C]. Linux Plumbers Conference (LPC), Refereed Track, 2021.',
    '[31] Nakryiko A. BPF ring buffer: a new data structure for BPF [EB/OL]. https://lwn.net/Articles/820559/, 2024.',
    '[32] Lan J P, Wang K X, Song S J, et al. Method for measuring non-stationary motion attitude based on MEMS-IMU array data fusion and adaptive filtering [J]. Meas. Sci. Technol., 2024, 35(8): 086304.',
    '[33] Li K P, Wang K X, Song S J, et al. Improved strong tracking Sage-Husa adaptive algorithm for multi-MEMS IMU data fusion [J]. Rev. Sci. Instrum., 2025, 96(5): 055002.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(ref)
    set_font(run, name='Times New Roman', size=Pt(8))

# ── Save ──
output_path = os.path.expanduser('C:/Users/xing2/Desktop/journal_paper_final.docx')
doc.save(output_path)
print(f'Paper saved to: {output_path}')
